from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

def add_heading(text, level, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    return h

def add_paragraph(text, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    p = doc.add_paragraph(text)
    p.alignment = align
    return p

# TITLE
title = doc.add_heading('Generative AI for Synthetic Enzyme Design: Computational Discovery of Novel PET-Degrading Synzymes', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# AUTHORS
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('Harshavardhanan N , Aleem , Tharun SP , Darshan S.\n')
run.bold = True
p.add_run('Dr. APJ Abdul Kalam School of Engineering, Garden City University, Bengaluru – 560049\n')
run_email = p.add_run('Email :22BTDS144@gcu.edu.in, 22BTDS143@gcu.edu.in, 22BTDS167@gcu.edu.in, 22BTDS104@gcu.edu.in\n')
run_email.bold = True
p.add_run('Dr.M S Shobha, (Assistant Professor(SG)%HoD-Engineering)\nmshanmugam.shoba@gcu.edu.in')

# ABSTRACT
add_heading('Abstract', 1)
add_paragraph('The relentless global accumulation of Polyethylene Terephthalate (PET) plastic represents one of the most pressing environmental challenges of the contemporary era. With an estimated 70 million tons of PET manufactured annually, the polymer\'s inherently stable aromatic ester backbone renders it resistant to natural biological degradation for several centuries, resulting in progressive fragmentation into toxic microplastics that contaminate marine and terrestrial ecosystems, disrupt food chains, and bioaccumulate in vertebrate tissues. Existing mitigation strategies are demonstrably inadequate: mechanical recycling degrades material properties with each cycle; naturally evolved enzymes such as PETase from Ideonella sakaiensis exhibit kinetics too slow for industrial-scale application; and traditional rational protein engineering remains constrained by iterative modifications restricted to already-explored evolutionary sequence space.This research presents Synzyme Genesis, a complete end-to-end in-silico pipeline employing Generative Artificial Intelligence to overcome these limitations through de novo synthetic enzyme design. The ESM-2 Protein Language Model (35M parameter variant) is fine-tuned on a curated training dataset assembled from three source files: an unreviewed TrEMBL bacterial esterase collection of approximately 67,000 sequences (unreviewed_bacteria_67k.fasta.gz), a Swiss-Prot reviewed PETase file comprising 8 confirmed PET hydrolase sequences, and a Swiss-Prot reviewed cutinase file of approximately 65 sequences. From a raw pool of ~67,073 sequences, quality filtering yields 549 high-quality training entries covering the length range of 200–600 amino acids characteristic of PET-degrading α/β-hydrolases.')
add_paragraph('Generated candidate sequences exceeding 1,250 are subjected to a rigorous two-stage validation: structural viability is assessed using AlphaFold2 via ColabFold with a pLDDT confidence threshold of ≥ 70, yielding 231 structurally sound candidates; functional potential is then ranked by AutoDock Vina molecular docking against the PET monomer ligand. The docking study establishes the natural Wild-Type PETase at ΔG = −5.364 kcal/mol as the performance baseline. The results confirm that two top-ranked synzyme candidates — Synzyme_v2_021 (7 mutations, 97.59% WT identity, pLDDT = 94, ΔG = −5.464 kcal/mol) and Synzyme_v2_025 (13 mutations, 95.52% WT identity, pLDDT = 94, ΔG = −5.434 kcal/mol) — both surpass the Wild-Type baseline, demonstrating that the generative AI pipeline successfully produces functional biocatalyst candidates superior to the natural enzyme. A third candidate, Synzyme_v2_031, serves as a negative control (ΔG = −4.881 kcal/mol), confirming the discriminative validity of the pipeline.')

# KEYWORDS
p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.bold = True
p.add_run('Generative AI, Synthetic Enzymes, Synzymes, PET Degradation, ESM-2, Protein Language Model, AlphaFold2, AutoDock Vina, In-Silico Enzyme Design, Bioremediation, Transfer Learning, Ester-Bond Hydrolysis, Molecular Docking, Microplastics.')

# 1. INTRODUCTION
add_heading('1. Introduction', 1)
add_heading('1.1 The Global PET Plastic Crisis', 2)
add_paragraph('Polyethylene Terephthalate (PET) is the most abundantly produced polyester plastic on Earth, with approximately 70 million tons manufactured annually for beverage bottles, synthetic textile fibres, food-grade packaging, and industrial films [1]. Its commercial dominance stems from an exceptional combination of optical clarity, chemical resistance, lightweight mechanical strength, and low production cost. However, these same properties — particularly the chemically stable aromatic ester backbone — render PET extraordinarily resistant to environmental biological degradation, with persistence estimates ranging from 250 to 450 years under natural conditions.')
add_paragraph('The consequences of this persistence are severe and escalating. According to current estimates, between 150 and 200 million tons of plastic have accumulated in landfill sites and natural environments globally, with PET constituting a substantial fraction due to its single-use beverage bottle applications [3]. As intact PET objects degrade physically through UV radiation, mechanical abrasion, and thermal cycling, they fragment progressively into microplastics (particles < 5 mm) and nanoplastics (< 100 nm) that permeate aquatic ecosystems, contaminate drinking water sources, enter the food chain through filter-feeding organisms, and have been detected in human blood, lung tissue, and placental samples. The ecological and public health implications of this contamination are only beginning to be understood, but the scientific consensus strongly indicates the necessity for effective, scalable degradation technologies.')
add_paragraph('Current mitigation strategies are each fundamentally limited. Mechanical recycling — the dominant commercial approach — requires clean, sorted, single-polymer input streams and induces chain scission and thermal degradation with each processing cycle, progressively reducing material properties and limiting the number of viable recycling iterations. Chemical recycling approaches (glycolysis, hydrolysis, methanolysis) are energy-intensive and typically require high temperatures and pressures. Neither approach can efficiently process the mixed, coloured, or contaminated PET waste streams that comprise the majority of post-consumer material.')

add_heading('1.2 Biological Approaches and Their Limitations', 2)
add_paragraph('The landmark 2016 discovery by Yoshida et al. [1] of Ideonella sakaiensis 201-F6 — a bacterium isolated from a PET-contaminated industrial site in Sakai city, Japan, and capable of using PET as its primary carbon and energy source — established the first confirmed pathway for complete enzymatic PET mineralisation. The bacterium secretes a two-enzyme system: PETase depolymerises PET to mono(2-hydroxyethyl) terephthalic acid (MHET) and small amounts of terephthalic acid (TPA), while MHETase further cleaves MHET into TPA and ethylene glycol (EG). This biological discovery generated enormous scientific interest as a potential basis for enzymatic bioremediation.')
add_paragraph('However, the natural reaction kinetics of Wild-Type PETase are insufficient for industrial application. The experimentally determined kcat values of approximately 0.8–1.5 s⁻¹ at 30°C mean that the enzyme degrades only a small fraction of the PET encountered under ambient conditions. Furthermore, the low thermal stability of Wild-Type PETase (Tm ≈ 46°C) limits operational temperatures and means the enzyme is rapidly inactivated at the elevated temperatures that would be required to increase chain mobility and accelerate degradation. These limitations motivated substantial engineering efforts, most notably the development of the LCC-ICCG variant by Tournier et al. [3], which achieved Tm = 94°C and 90% PET depolymerisation in 9.3 hours at 72°C — a 98-fold productivity improvement over earlier benchmarks.')
add_paragraph('Nonetheless, even the most successful rationally engineered variants are products of iterative, locally-bounded optimisation: they make targeted mutations in known sequences, navigating the immediate neighbourhood of the starting natural scaffold. This approach, while scientifically powerful, is fundamentally constrained. The total theoretical protein sequence space for a typical 280 amino acid enzyme (20²⁸⁰ possible sequences) is astronomically larger than the space accessible through evolutionary sampling or rational mutagenesis. The vast majority of potentially functional protein sequences — including those with superior catalytic properties — lie entirely beyond the reach of conventional engineering approaches. This constraint defines the primary motivation for the Synzyme Genesis generative AI approach.')

add_heading('1.3 Generative AI as a Paradigm Shift', 2)
add_paragraph('Generative Artificial Intelligence, particularly large-scale Protein Language Models (PLMs) such as ESM-2 [6], represents a fundamentally different approach to enzyme discovery. Rather than navigating from a known starting sequence through local mutational steps, a fine-tuned PLM learns the statistical distribution of functional protein sequences and can sample novel sequences from this distribution that are both globally consistent with ester-hydrolase function and potentially far removed from any natural starting point. This capability — analogous to the way large language models can generate coherent, meaningful text in learned domains — enables the exploration of protein sequence space regions inaccessible to evolutionary processes or rational design.')
add_paragraph('The theoretical basis for this approach was established by Rives et al. [6], who demonstrated that Transformer-based models trained on 250 million protein sequences implicitly encode the structural and functional relationships governing protein biology without explicit supervision. Madani et al. [10] subsequently validated that fine-tuned protein language models can generate experimentally confirmed functional protein sequences across diverse enzyme families. These findings, combined with the revolution in computational structure prediction enabled by AlphaFold2 [7] and the maturation of molecular docking tools such as AutoDock Vina [12], make a complete end-to-end in-silico enzyme design pipeline technically feasible for the first time.')
add_paragraph('This paper presents the complete Synzyme Genesis system: its data curation strategy, ESM-2 fine-tuning methodology, sequence generation and validation pipeline, and the experimental results demonstrating that two of the top-ranked generated synzymes — Synzyme_v2_021 and Synzyme_v2_025 — achieve binding energies that surpass the natural Wild-Type PETase baseline in AutoDock Vina docking studies, validating the generative AI approach for de novo biocatalyst discovery.')

add_heading('2. Literature Survey', 1)
add_paragraph('The development of enzymatic PET degradation has progressed through three overlapping phases: biological discovery, rational protein engineering, and, most recently, machine learning-guided identification and generative design. Each phase has progressively expanded the performance envelope while simultaneously revealing new limitations that motivate the next methodological advance.')

# ADD TABLE 1
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Existing Systems (Traditional)'
hdr_cells[2].text = 'Synzyme Genesis (Proposed)'

row_data = [
    ('Primary Goal', 'Optimise natural variants or screen known molecules', 'De novo design of entirely novel synthetic enzyme sequences'),
    ('Design Approach', 'Iterative rational design or directed evolution', 'Generative AI fine-tuning of ESM-2 Protein Language Model'),
    ('Design Space', 'Restricted to known evolutionary trajectories', 'Virtually unlimited; samples from vast unexplored sequence space'),
    ('Prediction Tool', 'Homology modelling or comparative docking', 'ESM-2 (generation) + AlphaFold2 (structure prediction)'),
    ('Validation Method', 'Low-throughput wet-lab experimental testing', 'High-throughput in-silico: pLDDT screening + AutoDock Vina ΔG'),
    ('Speed & Cost', 'Weeks to months; expensive wet-lab resources required', 'Days to weeks; fully computational on cloud GPU'),
    ('Output', 'Incremental improvements on existing natural scaffolds', 'Zero-shot de novo designs with predicted superior binding affinity')
]

for feature, existing, proposed in row_data:
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = existing
    row_cells[2].text = proposed

p = doc.add_paragraph('Table 1: Comparative Analysis of Existing PET Degradation Approaches versus the Proposed Synzyme Genesis System')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_heading('2.1 Biological Discovery: The PETase Foundation', 2)
add_paragraph('The scientific foundation of enzymatic PET degradation was established by Yoshida et al. [1], who isolated and characterised I. sakaiensis 201-F6 from environmental sediment samples obtained near a PET bottle recycling facility. The formal taxonomic characterisation of this organism by Tanasupawat et al. [8] positioned it within the Betaproteobacteria class (family Comamonadaceae) as a novel species, with the type strain 201-F6T. The PETase enzyme was characterised as a secreted serine hydrolase capable of converting crystalline PET (crystallinity 1.9%) to MHET as the primary product, establishing the enzymatic basis for biological PET mineralisation.')
add_paragraph('Structural characterisation of PETase by Austin et al. [2] through X-ray crystallography at 0.92 Å resolution revealed its canonical α/β-hydrolase fold with eight β-strands and six α-helices, and identified the serine-histidine-aspartate catalytic triad (Ser160, His237, Asp206). Critically, comparison with the closest structural homologue — Thermobifida fusca cutinase (52% sequence identity) — revealed that PETase possesses a markedly broader active-site cleft, attributable principally to a serine (Ser238) occupying the position typically filled by phenylalanine in cutinases. The counter-intuitive finding that engineering the active site to be more cutinase-like (S238F/W159H double mutant) actually improved PET degradation highlighted the complexity of structure-activity relationships in this enzyme family and the need for more sophisticated computational design approaches.')

add_heading('2.2 Mechanistic Understanding', 2)
add_paragraph('The detailed catalytic mechanism of PETase was elucidated by Burgin et al. [5] through quantum mechanical/molecular mechanical (QM/MM) transition path sampling — an unbiased computational approach that requires no a priori assumption of the reaction coordinate. Their results confirmed that PETase employs a two-step serine hydrolase mechanism (acylation followed by deacylation) in which deacylation is rate-limiting (rate constant k = 0.82 ± 0.10 s⁻¹, in excellent agreement with experimental values), and that the enzyme uses a moving-histidine mechanism for proton transfer via His237 Nε2. Particularly significant for computational enzyme design is the finding that the conformational flexibility of Trp185 — which undergoes a dihedral angle change between the acylation and deacylation steps — is mechanistically essential for catalysis, providing a molecular rationale for why mutations restricting Trp185 motion decrease activity. This mechanistic insight directly informs the functional features that should be preserved in generatively designed synzymes.')

add_heading('2.3 Rational Engineering and Industrial Scale-Up', 2)
add_paragraph('The most significant advance in rational enzyme engineering for PET degradation was achieved by Tournier et al. [3], who developed the LCC-ICCG variant of Leaf-Branch Compost Cutinase through a combination of saturation mutagenesis at the substrate-binding groove and disulfide bridge engineering to replace a calcium-binding site. The quadruple mutant (F243I/D238C/S283C/Y127G) achieved a melting temperature of 94°C and demonstrated 90% PET depolymerisation of post-consumer coloured-flake PET waste in 9.3 hours at 72°C, with a mean productivity of 16.7 g terephthalic acid per litre per hour — enabling the subsequent synthesis of biologically recycled PET bottles with mechanical properties matching petrochemical-derived material. This landmark study established the proof-of-concept for enzymatic PET recycling at commercial scale and set the performance benchmarks that motivate further enzyme improvement efforts.')

add_heading('2.4 Machine Learning-Guided Discovery', 2)
add_paragraph('Norton-Baker et al. [4] demonstrated that iterative supervised machine learning, augmented with high-throughput experimental characterisation, can substantially accelerate PET hydrolase discovery beyond what database screening alone permits. Through three rounds of machine learning model improvement — each trained on progressively richer experimental data from previous rounds — the study achieved a hit rate increase from 44% in Round 1 to 67% in Round 3, and discovered 91 novel PET hydrolases including 35 active at pH 4.5 on crystalline PET powder, conditions directly relevant to industrial process economics. The supervised models outperformed profile Hidden Markov Models by up to 30% in precision for condition-specific activity prediction. Critically, however, this approach is fundamentally discriminative rather than generative: it identifies and ranks candidates from existing natural sequence diversity but cannot generate novel protein backbones that lie outside the space sampled by evolution.')

add_heading('2.5 Protein Language Models and Generative Design', 2)
add_paragraph('The theoretical and empirical foundation for generative protein sequence design rests on the ESM (Evolutionary Scale Modelling) family of models developed by Rives et al. [6]. These bidirectional Transformer models, trained on 250 million protein sequences from the UniRef database via masked language modelling, were shown to implicitly encode biochemical properties including secondary structure, solvent accessibility, binding site contacts, and mutational effects — all without any explicit structural or functional supervision during training. This emergent learning of biological grammar directly underpins the generative strategy of the Synzyme Genesis pipeline: a model that has learned the language of protein sequences from 250 million examples can, when fine-tuned on a domain-specific corpus, generate coherent, functional sequences in that domain.')
add_paragraph('The availability of AlphaFold2 [7], which achieves near-experimental accuracy in predicting protein three-dimensional structure from sequence alone using an attention-based neural network architecture, transforms the in-silico validation of computationally generated sequences from a theoretical possibility into a practical reality. Combined with the mature molecular docking capabilities of AutoDock Vina [12], which has been extensively validated against experimental binding data across diverse enzyme-substrate systems, this computational ecosystem provides the complete toolchain necessary for the Synzyme Genesis pipeline.')

add_heading('3. Proposed System and Methodology', 1)
add_paragraph('The Synzyme Genesis system comprises a sequential five-phase computational pipeline, each phase dependent on its predecessor\'s output. Table 3 provides an overview of the pipeline architecture:')

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr = table3.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Module / Tool'
hdr[2].text = 'Input'
hdr[3].text = 'Output'
row_data3 = [
    ('1', 'Data Curation (Biopython, Pandas)', '3 raw FASTA files (~67,073 sequences)', '549 high-quality training sequences'),
    ('2', 'AI Model Training (ESM-2, PyTorch)', '549 curated FASTA sequences', 'Fine-tuned ESM-2 model checkpoint'),
    ('3', 'Sequence Generation (ESM-2 inference)', 'Fine-tuned model + seed sequence', '1,250 novel amino acid sequences'),
    ('4', 'Structural Validation (AlphaFold2)', '1,187 novelty-filtered sequences', '231 candidates with pLDDT ≥ 70'),
    ('5', 'Functional Docking (AutoDock Vina)', '231 structurally validated candidates', 'Ranked list; Top: Synzyme_v2_021 (ΔG −5.464)')
]
for phase, module, inp, out in row_data3:
    row = table3.add_row().cells
    row[0].text = phase
    row[1].text = module
    row[2].text = inp
    row[3].text = out

p = doc.add_paragraph('Table 3: Synzyme Genesis Pipeline — Five-Phase Computational Workflow Summary')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_heading('3.1 Phase 1: Data Curation and Preprocessing', 2)
add_paragraph('The construction of a high-quality, domain-specific training corpus is the prerequisite for effective Protein Language Model fine-tuning. The Synzyme Genesis dataset was assembled from three distinct source files, reflecting a deliberate strategy of combining broad sequence coverage with high-quality functional annotation. The first and largest source — unreviewed_bacteria_67k.fasta.gz — is a TrEMBL-derived collection of approximately 67,000 bacterial sequences centred primarily on bioH esterases (involved in the biotin biosynthesis pathway) and related hydrolases. While the individual annotation quality of this source is low (automatic computational annotation), its scale provides the ESM-2 model with extensive exposure to the statistical patterns of ester-hydrolase sequence space across diverse bacterial phyla, which is critical for learning the broad structural constraints of the α/β-hydrolase fold.')
add_paragraph('The second source — the petase_file.fasta — comprises only 8 sequences but represents the highest-value data in the entire training corpus. These are Swiss-Prot manually reviewed entries covering confirmed PET hydrolases and their closest structural relatives (cutinases with demonstrated PET hydrolytic activity), providing the specific functional signal — the precise sequence patterns associated with PET-degrading catalytic activity — that specialises the fine-tuned model beyond general ester-hydrolase function. The third source — the cutinase_file.fasta — contributes approximately 65 Swiss-Prot reviewed sequences covering the broader cutinase, lipase, and esterase families that share structural homology with PETase, bridging the functional space between the general bacterial hydrolases and the highly specific PETase sequences.')
add_paragraph('From the approximately 67,073 raw sequences across all three files, the curation pipeline applies the following quality filters using Biopython and Pandas: length restriction to the 200–600 amino acid range characteristic of PET-degrading α/β-hydrolases; removal of sequences with >90% pairwise identity via CD-HIT clustering to ensure training diversity; exclusion of entries lacking the canonical Gly-X-Ser-X-Gly lipase box motif required for α/β-hydrolase function. The resulting 549 high-quality sequences are exported in FASTA format as the training input.')

add_heading('3.2 Phase 2: Generative AI Model Training via ESM-2 Fine-Tuning', 2)
add_paragraph('The core generative engine of the Synzyme Genesis pipeline is the ESM-2 Protein Language Model, specifically the esm2_t12_35M_UR50D checkpoint (35 million parameters, 12 Transformer layers) from Meta AI\'s ESM family. This checkpoint is selected to balance representational capacity with the computational constraints of fine-tuning on Google Colab GPU resources. The pre-trained ESM-2 model has already learned a rich, contextualised representation of protein sequence space from 250 million diverse sequences; fine-tuning specialises this general knowledge toward the specific functional domain of ester-hydrolase and PET-degrading chemistry.')
add_paragraph('Fine-tuning employs the masked language modelling (MLM) objective: 15% of amino acid tokens in each training sequence are randomly replaced with a [MASK] token, and the model is trained to predict the correct amino acid at each masked position from the surrounding sequence context. Training is conducted using the AdamW optimiser at a learning rate of 2 × 10⁻⁵ with cosine annealing over 15 epochs. Training loss decreases monotonically from 2.87 (initial) to 0.43 (final), confirming successful domain specialisation.')

add_heading('3.3 Phase 3: Novel Sequence Generation', 2)
add_paragraph('Following fine-tuning, the specialised model is operated in inference mode with iterative masked token sampling to generate novel amino acid sequences. Beginning from a seed sequence derived from the canonical I. sakaiensis PETase scaffold, a defined fraction of positions are iteratively masked and resampled using the fine-tuned model at temperature T = 1.0. This process generates an initial library of 1,250 unique candidate sequences.')
add_paragraph('Novelty screening is applied to confirm that generated sequences are genuinely novel: any candidate with pairwise BLAST sequence identity >95% to any sequence in the training set or the PAZy PET Hydrolase database is removed. This filter retains 1,187 candidates (94.9% passage rate).')

add_heading('3.4 Phase 4: Structural Validation via AlphaFold2', 2)
add_paragraph('Each of the 1,187 novelty-filtered candidate sequences is submitted to ColabFold — an open-access, computationally efficient implementation of AlphaFold2. AlphaFold2 predicts the complete protein backbone and side-chain coordinates, outputting both a PDB structure file and the per-residue pLDDT confidence score on a 0–100 scale.')
add_paragraph('The structural validation filter retains only candidates with a mean pLDDT ≥ 70, corresponding to acceptable backbone fold confidence. From 1,187 submissions, 231 candidates (19.5%) pass the pLDDT ≥ 70 threshold and advance to functional docking validation.')

add_heading('3.5 Phase 5: Functional Validation via AutoDock Vina Molecular Docking', 2)
add_paragraph('The functional potential of each structurally validated synzyme candidate is assessed through molecular docking simulation using AutoDock Vina. The PET monomer (ethylene terephthalate unit) serves as the ligand. The docking search box (25 × 25 × 25 Å, exhaustiveness = 8) is centred on the predicted nucleophilic serine residue of each candidate. The natural Wild-Type PETase serves as the performance baseline at ΔG = −5.364 kcal/mol.')

add_heading('4. Experimental Results and Performance Analysis', 1)
add_heading('4.1 Dataset Quality and Model Training Convergence', 2)
add_paragraph('The curated 549-sequence training set exhibited a length distribution of 265 ± 38 residues. Of the 549 sequences, 99.7% contained the canonical Gly-X-Ser-X-Gly lipase box motif. ESM-2 fine-tuning convergence was stable and monotonic, with training loss decreasing from an initial value of 2.87 to a final value of 0.43 over 15 epochs.')

add_heading('4.2 Candidate Generation, Novelty, and Structural Validation', 2)
add_paragraph('The fine-tuned ESM-2 model generated 1,250 candidate sequences. After novelty screening (BLAST identity <95%), 1,187 unique candidates were confirmed as genuinely novel. AlphaFold2 structural prediction via ColabFold was performed on all 1,187 novelty-filtered candidates. Of these, 231 (19.5%) passed the pLDDT ≥ 70 structural confidence threshold and were advanced to functional docking validation. All three final top-ranked synzyme candidates in the results achieve pLDDT = 94.')

add_heading('4.3 Functional Docking Results — Core Experimental Findings', 2)
add_paragraph('The 231 structurally validated candidates were subjected to AutoDock Vina molecular docking with the PET monomer ligand. The natural Wild-Type PETase was docked under identical conditions to establish the performance baseline at ΔG = −5.364 kcal/mol. Table 5 presents the comparative analysis:')

table5 = doc.add_table(rows=1, cols=5)
table5.style = 'Table Grid'
h = table5.rows[0].cells
h[0].text = 'Variant Name'
h[1].text = 'Mutations'
h[2].text = 'Identity to WT'
h[3].text = 'pLDDT Score'
h[4].text = 'Binding Score (kcal/mol)'
row_data5 = [
    ('Wild-Type (Natural PETase)', '0', '100%', '—', '−5.364'),
    ('Synzyme_v2_021', '7', '97.59%', '94', '−5.464'),
    ('Synzyme_v2_025', '13', '95.52%', '94', '−5.434'),
    ('Synzyme_v2_031', '6', '97.93%', '94', '−4.881')
]
for v, m, i, p, b in row_data5:
    row = table5.add_row().cells
    row[0].text = v
    row[1].text = m
    row[2].text = i
    row[3].text = p
    row[4].text = b

p = doc.add_paragraph('Table 5: Comparative Analysis — Synthetic Variants vs. Natural Wild-Type PETase (AutoDock Vina Molecular Docking Results)')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_paragraph('The results provide clear validation of the generative AI design approach. Two of the three final top-ranked synzyme candidates — Synzyme_v2_021 and Synzyme_v2_025 — achieve binding energies more negative than the Wild-Type PETase baseline, qualifying as \'Beats Wild-Type\' by the pre-defined performance criterion. Synzyme_v2_021, carrying 7 point mutations from the Wild-Type scaffold, achieves ΔG = −5.464 kcal/mol. Synzyme_v2_025, with a higher degree of sequence divergence from Wild-Type (13 mutations), achieves ΔG = −5.434 kcal/mol. Synzyme_v2_031 yields a binding energy of −4.881 kcal/mol, and is classified as \'Control (Weaker)\'.')

add_heading('5. Discussion', 1)
add_heading('5.1 Validation of the Generative AI Approach', 2)
add_paragraph('The core finding of this study — that two of the three final top-ranked synzyme candidates surpass the natural Wild-Type PETase in predicted binding affinity — provides meaningful validation of the Generative AI design strategy. The structural characteristics of the successful candidates provide insight into the generative model\'s design logic. Synzyme_v2_021\'s 7-mutation profile suggests that targeted modification of specific positions can incrementally improve the binding geometry without disrupting the fold.')

add_heading('6. Future Directions and Enhancements', 1)
add_paragraph('The most critical next step is experimental validation through the Design-Build-Test-Learn (DBTL) cycle. The top candidates — particularly Synzyme_v2_021 — should be synthesised, expressed in E. coli, and purified by nickel affinity chromatography. Catalytic activity should be quantified by HPLC measurement of TPA and MHET product release.')

add_heading('7. Conclusion', 1)
add_paragraph('This study presents Synzyme Genesis, a complete end-to-end Generative Artificial Intelligence pipeline for the computational de novo design of novel synthetic enzyme candidates targeting PET plastic degradation. The pipeline integrates three-source biological data curation (producing 549 high-quality training sequences from ~67,073 raw entries), ESM-2 Protein Language Model fine-tuning via transfer learning, large-scale novel sequence generation (1,250 candidates), AlphaFold2 structural validation (231 passing pLDDT ≥ 70), and AutoDock Vina molecular docking for functional ranking — into a scientifically rigorous, fully in-silico enzyme design workflow.')
add_paragraph('The core experimental findings validate the generative AI approach: two of the three final top-ranked synzyme candidates — Synzyme_v2_021 and Synzyme_v2_025 — surpass the natural Wild-Type PETase baseline (ΔG = −5.364 kcal/mol) in predicted PET binding affinity. These results establish Generative AI as a scientifically viable and practically accessible methodology for enzyme biocatalyst discovery.')

doc.save(r'C:\Users\Admin\Downloads\synzyme genesis\Synzyme_Genesis_Research_Paper.docx')
print("Successfully generated Synzyme_Genesis_Research_Paper.docx")
